# -*- coding: utf-8 -*-
"""Evidence Factory core logic v3 (KNHANES + NHANES)
- Any variable can be chosen as exposure/outcome (continuous outcome = linear beta, binary = logistic OR)
- Smoking/alcohol grouped as never vs past/current, survey-weighted
- Covariates auto-selected once exposure and outcome are chosen
All statistics are deterministic (R survey). """
import pyreadstat, pandas as pd, numpy as np, glob, json, subprocess, os

# -- Unified variable registry (type: c=continuous, b=binary) --------
VARS={
 "age":("Age, years","c"),"bmi":("Body mass index, kg/m2","c"),"wc":("Waist circumference, cm","c"),
 "sbp":("Systolic blood pressure, mmHg","c"),"dbp":("Diastolic blood pressure, mmHg","c"),
 "glucose":("Fasting glucose, mg/dL","c"),"hba1c":("HbA1c, %","c"),"tg":("Triglycerides, mg/dL","c"),
 "tchol":("Total cholesterol, mg/dL","c"),"hdl":("HDL cholesterol, mg/dL","c"),
 "alt":("Alanine aminotransferase, IU/L","c"),"ast":("Aspartate aminotransferase, IU/L","c"),
 "ggt":("Gamma-glutamyl transferase, IU/L","c"),"insulin":("Fasting insulin, uIU/mL","c"),
 "fat_kg":("Total fat mass, kg","c"),"lean_kg":("Total lean mass, kg","c"),
 "asmm_kg":("Appendicular skeletal muscle mass, kg","c"),"bodyfat_pct":("Body fat, %","c"),
 "asm_pct":("Appendicular skeletal muscle, %","c"),"cap":("Controlled attenuation parameter, dB/m","c"),
 "men":("Male sex","b"),"smoking":("Ever smoking (past/current)","b"),
 "alcohol":("Ever alcohol (past/current)","b"),"htn":("Hypertension","b"),
 "dm":("Diabetes mellitus","b"),"mets":("Metabolic syndrome","b"),
 "dld":("Dyslipidemia","b"),"steatosis":("Hepatic steatosis","b")}
def lab(c): return VARS[c][0]
def typ(c): return VARS[c][1]

# Available variables per dataset (based on data currently on hand). Shared exposure/outcome list.
AVAIL={
 "KNHANES":[c for c in VARS if c not in ("ggt","cap")],
 "NHANES":[c for c in VARS if c not in ("hba1c","insulin","hdl","mets")]}
STEATOSIS_METHOD={"KNHANES":"NAFLD-LFS > -0.640 (index)","NHANES":"CAP >=248 dB/m (elastography)"}
DATASETS=["KNHANES","NHANES"]

def rscript_cmd():
    import shutil
    env=os.getenv("RSCRIPT")
    if env and os.path.exists(env): return env
    found=shutil.which("Rscript")
    if found: return found
    candidates=[]
    for base in [os.path.join(os.environ.get("ProgramFiles","C:\\Program Files"),"R"),
                 os.path.join(os.environ.get("ProgramFiles(x86)","C:\\Program Files (x86)"),"R")]:
        candidates += glob.glob(os.path.join(base,"R-*","bin","Rscript.exe"))
        candidates += glob.glob(os.path.join(base,"R-*","bin","x64","Rscript.exe"))
    candidates=[p for p in candidates if os.path.exists(p)]
    if candidates:
        return sorted(candidates)[-1]
    return "Rscript"
def r_env(workdir="."):
    env=os.environ.copy()
    if not env.get("R_LIBS_USER"):
        env["R_LIBS_USER"]=os.path.abspath(os.path.join(workdir,".Rlibs"))
    return env

# Definitional components of derived outcomes (prevents circularity + trajectory)
def determinants(dataset, outcome):
    base={"dm":["glucose","hba1c"],"htn":["sbp","dbp"],"mets":["wc","glucose","sbp","tg","hdl"],
          "dld":["tchol"]}
    if outcome=="steatosis": return ["cap"] if dataset=="NHANES" else ["insulin","ast","alt"]
    return [v for v in base.get(outcome,[]) if v in AVAIL.get(dataset,[])]

# Anthropometry/body-composition cluster (mutual adjustment within a cluster = overadjustment)
ADIPOSITY={"bmi","wc","fat_kg","lean_kg","asmm_kg","bodyfat_pct","asm_pct"}

# exposure/outcome -> auto covariate selection (avoids overadjustment)
def auto_covariates(dataset, exposures, outcomes):
    cov=["age","men"]                                   # age and sex always
    for c in ["smoking","alcohol"]:                     # lifestyle confounders
        if c in AVAIL[dataset]: cov.append(c)
    # BMI added as a confounder only when the exposure is not anthropometry/body-composition (avoid collinearity/mediation)
    if "bmi" in AVAIL[dataset] and not any(e in ADIPOSITY for e in exposures):
        cov.append("bmi")
    excl=set(exposures)|set(outcomes)
    for o in outcomes: excl|=set(determinants(dataset,o))   # exclude outcome definitional components
    if any(e in ADIPOSITY for e in exposures): excl|=ADIPOSITY  # if exposure is in the anthropometry cluster, exclude the whole cluster
    return [c for c in cov if c in AVAIL[dataset] and c not in excl]

DEFAULT_DEFS={
 "dm":{"FPG >=126 mg/dL":True,"HbA1c >=6.5%":True,"Medication":True,"Physician diagnosis":True},
 "htn":{"SBP >=140 or DBP >=90 mmHg":True,"Antihypertensive medication":True},
 "mets":{"Harmonized NCEP ATP III (3 of 5)":True},
 "dld":{"TC >=240 mg/dL":True,"HDL <40 mg/dL":True,"TG >=200 mg/dL":True},
 "steatosis":{"NAFLD-LFS >-0.640":True,"HSI >36 (alternative)":False},"pop":{"age_min":20}}

# -- Loaders ---------------------------------------------------------
def _find(dd,n):
    h=glob.glob(os.path.join(dd,"**",n),recursive=True); return h[0] if h else None
def _actual_usecols(path, desired, enc):
    if not desired: return None
    _,meta=pyreadstat.read_sas7bdat(path, metadataonly=True, encoding=enc)
    cols=list(meta.column_names)
    lower={c.lower():c for c in cols}
    return [c if c in cols else lower[c.lower()] for c in desired if c in cols or c.lower() in lower]
def _standardize_columns(df, desired):
    lower={c.lower():c for c in df.columns}
    ren={}
    for c in desired:
        actual=c if c in df.columns else lower.get(c.lower())
        if actual: ren[actual]=c
    return df.rename(columns=ren)
def _read_sas(path, usecols=None):
    encodings=[None,"cp949","euc-kr","latin1"]
    errors=[]
    for enc in encodings:
        try:
            actual=_actual_usecols(path, usecols, enc) if usecols else None
            df,meta=pyreadstat.read_sas7bdat(path, encoding=enc, usecols=actual)
            if usecols: df=_standardize_columns(df, usecols)
            return df,meta
        except Exception as e:
            errors.append(f"{enc or 'default'}: {e}")
    raise RuntimeError(f"Failed to read SAS file: {path}\n" + "\n".join(errors))
def required_data_message(dataset, data_dir, cycles):
    if dataset=="KNHANES":
        req=[]
        for y in cycles:
            req += [f"hn{y}_all.sas7bdat", f"hn{y}_dxa.sas7bdat"]
    else:
        mods=["demo","bmx","biopro","glu","trigly","tchol","lux","bpx","smq","alq","dxx"]
        req=[f"{m}_{cyc}.sas7bdat" for cyc in cycles for m in mods]
    found=glob.glob(os.path.join(data_dir,"**","*.sas7bdat"),recursive=True) if os.path.isdir(data_dir) else []
    sample=", ".join(os.path.basename(x) for x in found[:8]) if found else "none"
    return (f"Could not find {dataset} raw data. Current data folder: {data_dir}. "
            f"Required files (example): {', '.join(req[:6])}. "
            f"sas7bdat found in the current folder: {sample}")
def load_raw(dataset, data_dir, cycles):
    return _load_knhanes(data_dir,cycles) if dataset=="KNHANES" else _load_nhanes(data_dir,cycles)

def _load_knhanes(data_dir, years):
    ka=["id","kstrata","psu","wt_ex1","age","sex","HE_BMI","HE_wc","HE_glu","HE_HbA1c","HE_TG","HE_chol",
        "HE_HDL_st2","HE_ast","HE_alt","HE_insulin","HE_sbp","HE_dbp","DI1_2","DE1_31","DE1_32",
        "BS1_1","BD1","wt_tot","HE_HPdr","HE_DMdr","HE_HPdg","HE_DMdg","DE1_dg"]
    kx=["id","DW_WBT_FT","DW_WBT_LN","DW_WBT_MS","DW_Llg_LN","DW_Rlg_LN","DW_Lrm_LN","DW_Rrm_LN"]
    A,X=[],[]
    for y in years:
        fa,fx=_find(data_dir,f"hn{y}_all.sas7bdat"),_find(data_dir,f"hn{y}_dxa.sas7bdat")
        if not fa or not fx: continue
        a,_=_read_sas(fa, usecols=ka); x,_=_read_sas(fx, usecols=kx)
        a=a.rename(columns={c:"id" for c in a.columns if c.lower()=="id"})
        x=x.rename(columns={c:"id" for c in x.columns if c.lower()=="id"})
        A.append(a[[c for c in ka if c in a.columns]]); X.append(x[[c for c in kx if c in x.columns]])
    if not A:
        raise FileNotFoundError(required_data_message("KNHANES", data_dir, years))
    DF=pd.concat(A,ignore_index=True).merge(pd.concat(X,ignore_index=True),on="id",how="left")
    DF.attrs["dataset"]="KNHANES"; DF.attrs["n_cycles"]=len(A); return DF

def _load_nhanes(data_dir, cycles):
    spec={"demo":["SEQN","RIDAGEYR","RIAGENDR","SDMVSTRA","SDMVPSU","WTMEC2YR"],
          "bmx":["SEQN","BMXBMI","BMXWAIST"],"biopro":["SEQN","LBXSASSI","LBXSATSI","LBXSGTSI"],
          "glu":["SEQN","LBXGLU"],"trigly":["SEQN","LBXTR"],"tchol":["SEQN","LBXTC"],
          "hdl":["SEQN","LBDHDD"],"ghb":["SEQN","LBXGH"],"ins":["SEQN","LBXIN"],
          "lux":["SEQN","LUXCAPM","LUXSMED"],"bpx":["SEQN","BPXSY1","BPXDI1"],
          "smq":["SEQN","SMQ020","SMQ040"],"alq":["SEQN","ALQ111"],
          "dxx":["SEQN","DXDTOFAT","DXDTOLI","DXDTOTOT","DXXLALI","DXXRALI","DXXLLLI","DXXRLLI","DXXTRFAT"]}
    frames=[]
    for cyc in cycles:
        base=None
        for mod,want in spec.items():
            f=_find(data_dir,f"{mod}_{cyc}.sas7bdat")
            if not f: continue
            df,_=_read_sas(f, usecols=want); df=df[[c for c in want if c in df.columns]]
            base=df if base is None else base.merge(df,on="SEQN",how="left")
        if base is not None: frames.append(base)
    if not frames:
        raise FileNotFoundError(required_data_message("NHANES", data_dir, cycles))
    DF=pd.concat(frames,ignore_index=True); DF.attrs["dataset"]="NHANES"; DF.attrs["n_cycles"]=len(frames); return DF

# -- Preprocessing definitions -> standard cid columns (incl. smoking/alcohol grouping) --
def apply_definitions(DF, defs):
    ds=DF.attrs.get("dataset"); nc=DF.attrs.get("n_cycles",1)
    return _derive_knhanes(DF,defs,nc) if ds=="KNHANES" else _derive_nhanes(DF,defs,nc)

def _derive_knhanes(DF, de, nc):
    d=DF.copy()
    d["bmi"]=d.HE_BMI; d["wc"]=d.HE_wc; d["sbp"]=d.HE_sbp; d["dbp"]=d.HE_dbp; d["glucose"]=d.HE_glu
    d["hba1c"]=d.HE_HbA1c; d["tg"]=d.HE_TG; d["tchol"]=d.HE_chol; d["hdl"]=d.HE_HDL_st2
    d["alt"]=d.HE_alt; d["ast"]=d.HE_ast; d["insulin"]=d.HE_insulin; d["men"]=(d.sex==1).astype(int)
    limb=d[["DW_Llg_LN","DW_Rlg_LN","DW_Lrm_LN","DW_Rrm_LN"]].sum(axis=1)
    d["fat_kg"]=d.DW_WBT_FT/1000; d["lean_kg"]=d.DW_WBT_LN/1000; d["asmm_kg"]=limb/1000
    d["bodyfat_pct"]=d.DW_WBT_FT/d.DW_WBT_MS*100; d["asm_pct"]=limb/d.DW_WBT_MS*100
    # Smoking: never(BS1_1==3) vs ever(BS1_1 in 1,2)
    d["smoking"]=np.where(d.BS1_1==3,0,np.where(d.BS1_1.isin([1,2]),1,np.nan))
    # Alcohol: never(BD1==1) vs ever(BD1==2)
    d["alcohol"]=np.where(d.BD1==1,0,np.where(d.BD1==2,1,np.nan))
    def col(n): return d[n] if n in d.columns else pd.Series(np.nan,index=d.index)
    de31=col("DE1_31"); de32=col("DE1_32")
    hpdr=col("HE_HPdr"); dmdr=col("HE_DMdr"); hpdg=col("HE_HPdg"); dmdg=col("HE_DMdg")
    male=d.sex==1; female=d.sex==2
    # Hypertension: SBP>=140 or DBP>=90 or on antihypertensive medication
    htn=pd.Series(False,index=d.index)
    if de["htn"].get("SBP >=140 or DBP >=90 mmHg"): htn|=(d.sbp>=140)|(d.dbp>=90)
    if de["htn"].get("Antihypertensive medication"): htn|=(hpdr==1)
    d["htn"]=htn.astype(int)
    # Diabetes: FPG>=126 or HbA1c>=6.5 or medication or diagnosis (ADA/KNHANES standard)
    dm=pd.Series(False,index=d.index)
    if de["dm"].get("FPG >=126 mg/dL"): dm|=(d.glucose>=126)
    if de["dm"].get("HbA1c >=6.5%"): dm|=((d.hba1c>=6.5)&d.hba1c.notna())
    if de["dm"].get("Medication"): dm|=(dmdr==1)|(de31==1)|(de32==1)
    if de["dm"].get("Physician diagnosis"): dm|=(dmdg==1)
    d["dm"]=dm.astype(int)
    # Metabolic syndrome: harmonized NCEP ATP III -- 3 of 5 (abdominal obesity not required)
    c_wc=((male&(d.wc>=90))|(female&(d.wc>=85)))
    c_glu=(d.glucose>=100)|(dmdr==1)
    c_bp=((d.sbp>=130)|(d.dbp>=85))|(hpdr==1)
    c_tg=(d.tg>=150)
    c_hdl=((d.hdl<40)&male)|((d.hdl<50)&female)
    cnt=sum(x.astype("boolean").fillna(False).astype(int) for x in [c_wc,c_glu,c_bp,c_tg,c_hdl])
    d["mets"]=(cnt>=3).astype(int)
    # Dyslipidemia: TC>=240 or HDL<40 or TG>=200 (Korean lipid guideline)
    dld=pd.Series(False,index=d.index)
    if de["dld"].get("TC >=240 mg/dL"): dld|=(d.tchol>=240)
    if de["dld"].get("HDL <40 mg/dL"): dld|=(d.hdl<40)
    if de["dld"].get("TG >=200 mg/dL"): dld|=(d.tg>=200)
    d["dld"]=dld.astype(int)
    # Hepatic steatosis: NAFLD-LFS (default) or HSI>36
    if de.get("steatosis",{}).get("HSI >36 (alternative)"):
        hsi=8*(d.alt/d.ast)+d.bmi+np.where(female,2,0)+2*d.dm
        d["steatosis"]=(hsi>36).astype("Int64")
    else:
        lfs=-2.89+1.18*d.mets+0.9*d.dm+0.15*d.insulin+0.04*d.ast-0.94*(d.ast/d.alt)
        d["steatosis"]=(lfs>-0.640).astype("Int64")
    wcol="wt_ex1" if "wt_ex1" in d.columns else ("wt_tot" if "wt_tot" in d.columns else None)
    d["kstrata"]=d.kstrata; d["psu"]=d.psu
    d["wt_pool"]=(d[wcol]/max(nc,1)) if wcol else np.nan
    return d

def _derive_nhanes(DF, de, nc):
    d=DF.copy()
    d["age"]=d.RIDAGEYR; d["men"]=(d.RIAGENDR==1).astype(int)
    d["bmi"]=d.BMXBMI; d["wc"]=d.BMXWAIST; d["sbp"]=d.BPXSY1; d["dbp"]=d.BPXDI1
    d["glucose"]=d.LBXGLU; d["tg"]=d.LBXTR; d["tchol"]=d.LBXTC
    d["ast"]=d.LBXSASSI; d["alt"]=d.LBXSATSI; d["ggt"]=d.LBXSGTSI; d["cap"]=d.LUXCAPM
    limb=d[["DXXLALI","DXXRALI","DXXLLLI","DXXRLLI"]].sum(axis=1); tot=d.DXDTOFAT+d.DXDTOLI
    d["fat_kg"]=d.DXDTOFAT/1000; d["lean_kg"]=d.DXDTOLI/1000; d["asmm_kg"]=limb/1000
    d["bodyfat_pct"]=d.DXDTOFAT/tot*100; d["asm_pct"]=limb/tot*100
    # Smoking: never(SMQ020==2) vs ever(SMQ020==1)
    d["smoking"]=np.where(d.SMQ020==1,1,np.where(d.SMQ020==2,0,np.nan))
    # Alcohol: never(ALQ111==2) vs ever(ALQ111==1)
    d["alcohol"]=np.where(d.ALQ111==1,1,np.where(d.ALQ111==2,0,np.nan))
    d["htn"]=(((d.sbp>=140)|(d.dbp>=90)) if de["htn"].get("SBP >=140 or DBP >=90 mmHg") else False)
    d["htn"]=pd.Series(d["htn"],index=d.index).astype(int)
    d["dm"]=((d.glucose>=126) if de["dm"].get("FPG >=126 mg/dL") else False)
    d["dm"]=pd.Series(d["dm"],index=d.index).astype(int)
    dld=pd.Series(False,index=d.index)
    if de["dld"].get("TC >=240 mg/dL"): dld|=(d.tchol>=240)
    if de["dld"].get("TG >=200 mg/dL"): dld|=(d.tg>=200)
    d["dld"]=dld.astype(int)
    d["steatosis"]=(d.LUXCAPM>=248).astype("Int64")
    d["kstrata"]=d.SDMVSTRA; d["psu"]=d.SDMVPSU; d["wt_pool"]=d.WTMEC2YR/max(nc,1)
    return d

# -- Analytic data + exposure term (continuous=z-standardized, binary=0/1 as-is) --
def build_analytic(d, exposures, outcomes, covariates, age_min):
    keep=(d.age>=age_min)&(d.wt_pool>0)&d.wt_pool.notna()&d.kstrata.notna()&d.psu.notna()
    if "bodyfat_pct" in d: keep&=d.bodyfat_pct.notna()
    ana=d[keep].copy()
    for o in outcomes: ana=ana[ana[o].notna()]
    for e in exposures:
        if typ(e)=="c": ana[f"x_{e}"]=(ana[e]-ana[e].mean())/ana[e].std()
        else: ana[f"x_{e}"]=ana[e]
    return ana

def run_engine(ana, dataset, exposures, outcomes, covariates, workdir="."):
    cont_std=["age","bmi","wc","sbp","glucose","tg","alt","fat_kg","asmm_kg","bodyfat_pct","asm_pct"]
    cont=[c for c in dict.fromkeys(cont_std+[e for e in exposures if typ(e)=="c"]) if c in ana.columns]
    primary=outcomes[0]
    group=primary if typ(primary)=="b" else ""     # stratify if binary outcome, whole sample if continuous
    binv=[b for b in ["men","smoking","alcohol","htn","dm","mets","dld"]
          if b!=group and b in ana.columns and ana[b].notna().any()]
    cov=[c for c in covariates if c in ana.columns]
    pairs=[{"y":o,"x":f"x_{e}","otype":typ(o)} for e in exposures for o in outcomes
           if e!=o and e not in determinants(dataset,o)]
    labels={k:lab(k) for k in cont+binv}
    cfg={"group":group,"cont":cont,"bin":binv,"cov":cov,"labels":labels,"pairs":pairs}
    ana.to_csv(os.path.join(workdir,"engine_analytic.csv"),index=False)
    json.dump(cfg,open(os.path.join(workdir,"engine_config.json"),"w"),ensure_ascii=False)
    r=subprocess.run([rscript_cmd(),"engine.R"],cwd=workdir,capture_output=True,text=True,env=r_env(workdir))
    if r.returncode!=0: raise RuntimeError(r.stderr)
    return (pd.read_csv(os.path.join(workdir,"engine_table1.csv")),
            pd.read_csv(os.path.join(workdir,"engine_results.csv")), cfg, pairs)

def add_fdr(res):
    if len(res)==0: return res
    p=res.p.values; n=len(p); o=np.argsort(p)
    q=np.empty(n); q[o]=np.minimum.accumulate((p[o]*n/np.arange(1,n+1))[::-1])[::-1]
    res=res.copy(); res["q"]=np.clip(q,0,1); return res

def effect_str(r):
    return f"{r.est:.2f} ({r.ci_low:.2f}-{r.ci_high:.2f})"
def measure_name(m): return "OR" if m=="OR" else "beta"

# -- LLM (ollama) + fallback -----------------------------------------
def ollama_chat(prompt, model=None, url=None, fmt=None):
    import local_llm
    return local_llm.generate(prompt, model=model, url=url, fmt=fmt)
STYLE="Medical journal prose style. Do not use em dashes, en dashes, arrows, semicolons, colons, or (i)(ii) enumeration in the body. First/Second/Third are allowed. Do not change any numbers."

def gen_manuscript(t1,res,dataset,exposures,outcomes,covariates,defs,model,url,use_llm):
    res=add_fdr(res)
    exp_n=", ".join(lab(e) for e in exposures); out_n=", ".join(lab(o) for o in outcomes)
    cov_n=", ".join(lab(c) for c in covariates)
    src={"KNHANES":"Korea National Health and Nutrition Examination Survey",
         "NHANES":"National Health and Nutrition Examination Survey"}[dataset]
    fnd="; ".join(f"{lab(r.exposure[2:])} and {lab(r.outcome)} {measure_name(r.measure)} {r.est:.2f} "
                  f"(95% CI {r.ci_low:.2f}-{r.ci_high:.2f}, FDR q {r.q:.3f})" for _,r in res.iterrows())
    facts=(f"Data: {src}, survey-weighted. Exposures: {exp_n}. Outcomes: {out_n}. Covariates: {cov_n}. "
           f"Smoking and alcohol grouped as never versus past or current. Continuous exposures modeled per 1-SD. "
           f"Binary outcomes used survey logistic regression reporting odds ratios; continuous outcomes used "
           f"survey linear regression reporting beta coefficients. Findings: {fnd}.")
    if use_llm:
        try:
            def sec(nm,ins): return ollama_chat(f"{STYLE}\nWrite the {nm} paragraph of a research paper in English, based only on the facts below. Do not change any numbers.\nFacts: {facts}\n{ins}",model,url).strip()
            return {"Methods":sec("Methods","Describe data, survey design, variable grouping, and models."),
                    "Results":sec("Results","Report each association with its effect measure, CI, FDR q."),
                    "Discussion":sec("Discussion","Interpret, note hypothesis-generating nature and limitations."),
                    "_source":"llm"}
        except Exception: pass
    meth=(f"We analyzed {src} participants using complex-survey methods accounting for stratification, clustering, "
          f"and sampling weights. Smoking and alcohol were grouped as never versus past or current. {exp_n} served "
          f"as exposures and {out_n} as outcomes, adjusting for {cov_n}. Continuous exposures were modeled per 1-SD. "
          f"Binary outcomes used survey-weighted logistic regression, and continuous outcomes used survey-weighted "
          f"linear regression. Multiplicity was controlled with the Benjamini-Hochberg false discovery rate.")
    rl=[]
    for _,r in res.sort_values("q").iterrows():
        if r.measure=="OR":
            rl.append(f"{lab(r.exposure[2:])} was {'inversely' if r.est<1 else 'positively'} associated with "
                      f"{lab(r.outcome)} (odds ratio {r.est:.2f}, 95% CI {r.ci_low:.2f} to {r.ci_high:.2f}, FDR q {r.q:.3f}).")
        else:
            rl.append(f"{lab(r.exposure[2:])} was associated with a {'higher' if r.est>0 else 'lower'} {lab(r.outcome)} "
                      f"(beta {r.est:.2f}, 95% CI {r.ci_low:.2f} to {r.ci_high:.2f}, FDR q {r.q:.3f}).")
    disc=("These survey-weighted associations are hypothesis-generating and require external validation. "
          "The cross-sectional design precludes causal inference.")
    return {"Methods":meth,"Results":" ".join(rl),"Discussion":disc,"_source":"fallback"}

def build_docx(title, sections, t1, res, group_label=""):
    from docx import Document
    res=add_fdr(res); doc=Document(); doc.add_heading(title,0)
    for nm in ["Methods","Results","Discussion"]:
        doc.add_heading(nm,1); doc.add_paragraph(sections.get(nm,""))
    doc.add_heading("Table 1. Descriptive characteristics",1)
    cols=list(t1.columns); tb=doc.add_table(rows=1,cols=len(cols)); tb.style="Light Grid Accent 1"
    for j,c in enumerate(cols): tb.rows[0].cells[j].text=str(c)
    for _,row in t1.iterrows():
        cc=tb.add_row().cells
        for j,c in enumerate(cols): cc[j].text=str(row[c])
    doc.add_heading("Table 2. Adjusted associations",1)
    rc=["Exposure","Outcome","Measure","Estimate (95% CI)","FDR q"]; tb2=doc.add_table(rows=1,cols=5); tb2.style="Light Grid Accent 1"
    for j,c in enumerate(rc): tb2.rows[0].cells[j].text=c
    for _,r in res.sort_values("q").iterrows():
        cc=tb2.add_row().cells
        cc[0].text=lab(r.exposure[2:]); cc[1].text=lab(r.outcome); cc[2].text=measure_name(r.measure)
        cc[3].text=effect_str(r); cc[4].text=f"{r.q:.3f}"
    import io; buf=io.BytesIO(); doc.save(buf); return buf.getvalue()

def nl_to_config(question, dataset, model, url, use_llm):
    av=AVAIL[dataset]
    if use_llm:
        try:
            o=json.loads(ollama_chat(f"Convert the question into an analysis configuration. Available variables (shared exposure/outcome)={av}. "
                f"Return JSON only: {{'exposures':[],'outcomes':[]}}.\nQuestion: {question}",model,url,fmt="json"))
            e=[x for x in o.get("exposures",[]) if x in av]; oo=[x for x in o.get("outcomes",[]) if x in av]
            if e and oo: return {"exposures":e,"outcomes":oo,"covariates":auto_covariates(dataset,e,oo)}
        except Exception: pass
    q=question.lower(); kw={"fatty liver":"steatosis","steatosis":"steatosis","diabetes":"dm","hypertension":"htn",
        "metabolic syndrome":"mets","dyslipidemia":"dld","muscle":"asm_pct","body fat":"bodyfat_pct","bmi":"bmi",
        "obesity":"bmi","waist":"wc","smoking":"smoking","alcohol":"alcohol",
        "glucose":"glucose","triglyceride":"tg","cholesterol":"tchol"}
    hits=[v for k,v in kw.items() if k in q and v in av]
    e=[hits[0]] if hits else ["bodyfat_pct"]; oo=[hits[1]] if len(hits)>1 else ["dm" if "dm" in av else av[-1]]
    return {"exposures":e,"outcomes":oo,"covariates":auto_covariates(dataset,e,oo)}

# -- Trajectory ------------------------------------------------------
def trajectory(d, dataset, outcome, split_sex=False, bin_width=5, age_range=(20,80)):
    vs=[v for v in determinants(dataset,outcome) if v in d.columns and d[v].notna().any()]
    dd=d[(d.age>=age_range[0])&(d.age<=age_range[1])&d.wt_pool.notna()&(d.wt_pool>0)].copy()
    dd["agebin"]=(np.floor(dd.age/bin_width)*bin_width).astype(int)
    def wmean(sub,col):
        s=sub[[col,"wt_pool"]].dropna()
        return float(np.average(s[col].astype(float),weights=s.wt_pool)) if len(s) else np.nan
    rows=[]; grouped=dd.groupby(["agebin","men"]) if split_sex else dd.groupby("agebin")
    for kv,sub in grouped:
        if split_sex: ab,mn=kv; sx="Men" if mn==1 else "Women"
        else: ab=kv; sx="All"
        for v in vs: rows.append({"agebin":int(ab),"sex":sx,"variable":lab(v),"value":wmean(sub,v)})
        rows.append({"agebin":int(ab),"sex":sx,"variable":f"{lab(outcome)} prevalence, %","value":wmean(sub,outcome)*100})
    return pd.DataFrame(rows), vs

def plot_trajectory(df, outcome, dataset, split_sex=False):
    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
    panels=list(dict.fromkeys(df.variable)); n=len(panels); ncol=min(3,n) or 1; nrow=(n+ncol-1)//ncol
    fig,axes=plt.subplots(nrow,ncol,figsize=(4*ncol,3*nrow),squeeze=False)
    for i,pan in enumerate(panels):
        ax=axes[i//ncol][i%ncol]; sub=df[df.variable==pan]
        if split_sex:
            for sx,g in sub.groupby("sex"): ax.plot(g.agebin,g.value,marker="o",label=sx)
            ax.legend(fontsize=7)
        else: ax.plot(sub.agebin,sub.value,marker="o",color="#0F6E56")
        ax.set_title(pan,fontsize=9); ax.set_xlabel("Age, years",fontsize=8); ax.grid(alpha=0.3)
    for j in range(n,nrow*ncol): axes[j//ncol][j%ncol].axis("off")
    fig.suptitle(f"{lab(outcome)} determinants -- age trajectory ({dataset}, survey-weighted)",fontsize=11)
    fig.tight_layout(); return fig

def fig_to_b64(fig):
    import io,base64,matplotlib.pyplot as plt
    buf=io.BytesIO(); fig.savefig(buf,format="png",dpi=90,bbox_inches="tight"); plt.close(fig)
    return base64.b64encode(buf.getvalue()).decode()

# -- Static HTML report (measure-aware) ------------------------------
def build_html_report(dataset, config, t1, res, manuscript, traj=None):
    import html as _h
    res=add_fdr(res); rows=""
    for _,r in res.sort_values("q").iterrows():
        risk = (r.measure=="OR" and r.est>1) or (r.measure=="beta" and r.est>0)
        col="#A32D2D" if risk else "#0F6E56"; sig="background:#EAF3DE;" if r.q<0.05 else ""
        rows+=(f'<tr style="{sig}"><td>{_h.escape(lab(r.exposure[2:]))}</td><td>{_h.escape(lab(r.outcome))}</td>'
               f'<td>{measure_name(r.measure)}</td>'
               f'<td style="color:{col};text-align:right;font-weight:500">{effect_str(r)}</td>'
               f'<td style="text-align:right">{"<0.001" if r.q<0.001 else f"{r.q:.3f}"}</td></tr>')
    t1_html=t1.to_html(index=False,border=0)
    ms_html="".join(f'<h2>{k}</h2><p>{_h.escape(str(manuscript.get(k,"")))}</p>' for k in ["Methods","Results","Discussion"])
    traj_html="".join(f'<div><h3>{_h.escape(t)}</h3><img src="data:image/png;base64,{b}"></div>' for t,b in (traj or []))
    exp_n=", ".join(lab(e) for e in config.get("exp",[])); out_n=", ".join(lab(o) for o in config.get("out",[]))
    cov_n=", ".join(lab(c) for c in config.get("cov",[])); src=manuscript.get("_source","fallback")
    return f"""<!DOCTYPE html><html lang=en><head><meta charset=utf-8><title>Evidence Report -- {dataset}</title><style>
body{{{{font-family:system-ui,'Segoe UI',sans-serif;max-width:1000px;margin:20px auto;padding:0 18px;color:#2c2c2a;line-height:1.5}}}}
h1{{{{font-size:22px;margin-bottom:2px}}}} h2{{{{font-size:17px;border-bottom:2px solid #d3d1c7;padding-bottom:4px;margin-top:26px}}}}
.bar{{{{background:#F1EFE8;border-radius:10px;padding:12px 16px;font-size:13px;margin:12px 0}}}}
table{{{{border-collapse:collapse;width:100%;font-size:13px;margin-top:8px}}}} th,td{{{{border-bottom:1px solid #e7e7e2;padding:7px;text-align:left}}}}
th{{{{background:#fafaf7;border-bottom:2px solid #d3d1c7}}}} img{{{{max-width:100%;border:1px solid #eee;border-radius:6px}}}}</style></head><body>
<h1>Evidence Report -- {dataset}</h1>
<div style="color:#666;font-size:14px">Survey-weighted &middot; continuous outcome = linear beta, binary = logistic OR &middot; steatosis definition: {STEATOSIS_METHOD[dataset]}</div>
<div class=bar><b>Exposure</b>: {_h.escape(exp_n)} &nbsp;|&nbsp; <b>Outcome</b>: {_h.escape(out_n)} &nbsp;|&nbsp;
<b>Auto-adjusted</b>: {_h.escape(cov_n)} &nbsp;|&nbsp; Prose: {src}</div>
{ms_html}
<h2>Table 1. Descriptive characteristics</h2>{t1_html}
<h2>Table 2. Adjusted associations</h2>
<table><thead><tr><th>Exposure</th><th>Outcome</th><th>Measure</th><th style=text-align:right>Estimate (95% CI)</th><th style=text-align:right>FDR q</th></tr></thead><tbody>{rows}</tbody></table>
<h2>Outcome determinants -- age trajectory</h2>{traj_html or "<p>(none)</p>"}
<p style="color:#888;font-size:12px;margin-top:18px">Statistics are deterministic R survey design-weighted estimates. Prose is written from the computed numbers (no recomputation). All findings are hypothesis-generating.</p>
</body></html>"""
