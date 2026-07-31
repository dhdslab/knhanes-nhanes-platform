# -*- coding: utf-8 -*-
"""역학분석 Word 보고서 생성 (Table1+SMD, crude/adjusted OR, 전체위험, 하위군+P-interaction, forest, E-value)"""
import factory_core as fc
import pandas as pd, numpy as np, json, subprocess, os, io

def _wstat(sub,col):
    s=sub[[col,"wt_pool"]].dropna(); w=s.wt_pool.values; x=s[col].values.astype(float)
    if w.sum()==0: return np.nan,np.nan
    m=np.average(x,weights=w); return m,np.average((x-m)**2,weights=w)

def table1_smd(d,outcome,variables):
    g0=d[d[outcome]==0]; g1=d[d[outcome]==1]; rows=[]
    for v in variables:
        binary=(fc.typ(v)=="b") or (d[v].dropna().nunique()<=2)
        if binary:
            p0,_=_wstat(g0,v); p1,_=_wstat(g1,v); pa,_=_wstat(d,v); pb=(p0+p1)/2
            smd=abs(p1-p0)/np.sqrt(pb*(1-pb)) if 0<pb<1 else np.nan
            rows.append({"변수":fc.lab(v),"전체":f"{100*pa:.1f}%","outcome=0":f"{100*p0:.1f}%","outcome=1":f"{100*p1:.1f}%","SMD":f"{smd:.3f}"})
        else:
            m0,v0=_wstat(g0,v); m1,v1=_wstat(g1,v); ma,_=_wstat(d,v); sp=np.sqrt((v0+v1)/2)
            smd=abs(m1-m0)/sp if sp>0 else np.nan
            rows.append({"변수":fc.lab(v),"전체":f"{ma:.1f}","outcome=0":f"{m0:.1f}","outcome=1":f"{m1:.1f}","SMD":f"{smd:.3f}"})
    return pd.DataFrame(rows)

def evalue(orr,lo,hi):
    ev=lambda x:round((1/x if x<1 else x)+np.sqrt((1/x if x<1 else x)*((1/x if x<1 else x)-1)),2)
    return ev(orr), ev(hi if orr<1 else lo)

def _forest(adj,labels):
    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
    a=adj.dropna(subset=["OR"]).copy(); a["name"]=a.term.map(labels)
    fig,ax=plt.subplots(figsize=(6,0.5*len(a)+1))
    y=np.arange(len(a))
    ax.errorbar(a.OR,y,xerr=[a.OR-a.lo,a.hi-a.OR],fmt="o",color="#245",capsize=3)
    ax.axvline(1,color="#999",ls="--"); ax.set_yticks(y); ax.set_yticklabels(a.name,fontsize=8)
    ax.set_xscale("log"); ax.set_xlabel("Adjusted OR (95% CI)"); ax.invert_yaxis(); fig.tight_layout(); return fig

def build_epi_report(dataset, DF, outcome, main_exposure, covariates, subgroups, defs, model, url, use_llm, workdir="."):
    d=fc.apply_definitions(DF,defs); amin=defs["pop"]["age_min"]
    n_total=len(d); d_age=d[d.age>=amin]; n_age=len(d_age)
    # 연속형 → z표준화 항 생성
    def term(v):
        if fc.typ(v)=="c":
            col=f"z_{v}"; d[col]=(d[v]-d[v].mean())/d[v].std(); return col
        return v
    d["age50"]=(d.age>=50).astype(int)  # 하위군용 연령 이분
    ex_t=term(main_exposure); cov_t=[term(c) for c in covariates]
    labels={ex_t:fc.lab(main_exposure),**{term(c):fc.lab(c) for c in covariates}}
    # 인과분석용 이분 노출(연속이면 가중 중앙값 기준 이분), RCS용 연속항
    def wmedian(x,w):
        s=pd.DataFrame({"x":x,"w":w}).dropna().sort_values("x"); c=s.w.cumsum()
        return s.x[c>=s.w.sum()/2].iloc[0]
    if fc.typ(main_exposure)=="c":
        med=wmedian(d[main_exposure],d["wt_pool"]); d["xbin"]=(d[main_exposure]>med).astype("Int64")
        exposure_bin="xbin"; exposure_cont=ex_t; xbin_label=f"{fc.lab(main_exposure)} 상위(중앙값 초과)"
    else:
        exposure_bin=main_exposure; exposure_cont=""; xbin_label=fc.lab(main_exposure)
    need=list(dict.fromkeys([outcome,ex_t,exposure_bin]+cov_t+subgroups+["age","kstrata","psu","wt_pool"]))
    ana=d[(d.age>=amin)&(d.wt_pool>0)&d.wt_pool.notna()&d.kstrata.notna()&d.psu.notna()].dropna(subset=[outcome]).copy()
    n_ana=len(ana)
    ana[[c for c in need if c in ana.columns]].to_csv(os.path.join(workdir,"epi_analytic.csv"),index=False)
    json.dump({"outcome":outcome,"exposure":ex_t,"cov":cov_t,"subgroups":subgroups,
               "exposure_bin":exposure_bin,"exposure_cont":exposure_cont},
              open(os.path.join(workdir,"epi_config.json"),"w"),ensure_ascii=False)
    r=subprocess.run([fc.rscript_cmd(),"epi.R"],cwd=workdir,capture_output=True,text=True,env=fc.r_env(workdir))
    if r.returncode!=0: raise RuntimeError(r.stderr)
    ra=subprocess.run([fc.rscript_cmd(),"epi_adv.R"],cwd=workdir,capture_output=True,text=True,env=fc.r_env(workdir))
    if ra.returncode!=0: raise RuntimeError("epi_adv: "+ra.stderr)
    vif=pd.read_csv(os.path.join(workdir,"epi_vif.csv")); rcs=pd.read_csv(os.path.join(workdir,"epi_rcs.csv"))
    love=pd.read_csv(os.path.join(workdir,"epi_love.csv")); tab6=pd.read_csv(os.path.join(workdir,"epi_table6.csv"))
    crude=pd.read_csv(os.path.join(workdir,"epi_crude.csv")); adj=pd.read_csv(os.path.join(workdir,"epi_adj.csv"))
    risk=pd.read_csv(os.path.join(workdir,"epi_risk.csv")); sub=pd.read_csv(os.path.join(workdir,"epi_sub.csv"))
    t1=table1_smd(ana,outcome,[main_exposure]+covariates)
    me=adj[adj.term==ex_t].iloc[0]; evp,evc=evalue(me.OR,me.lo,me.hi)
    fig=_forest(adj,labels)
    # LLM 해석 (선택)
    facts=(f"Outcome {fc.lab(outcome)}. Main exposure {fc.lab(main_exposure)} adjusted OR {me.OR:.2f} "
           f"(95% CI {me.lo:.2f}-{me.hi:.2f}). E-value {evp} (CI {evc}). Weighted risk {risk.weighted_risk_pct[0]:.1f}%.")
    interp=""
    if use_llm:
        try: interp=fc.ollama_chat(f"{fc.STYLE}\n아래 역학분석 결과를 논문 Results 단락으로. 수치변경 금지.\n{facts}",model,url).strip()
        except Exception: interp=""
    if not interp:
        interp=(f"In survey-weighted logistic regression, {fc.lab(main_exposure)} showed an adjusted odds ratio of "
                f"{me.OR:.2f} (95% CI {me.lo:.2f} to {me.hi:.2f}) for {fc.lab(outcome)}. The E-value was {evp}, "
                f"indicating the minimum strength of unmeasured confounding on the odds-ratio scale that could explain the estimate.")
    # ── docx ──
    from docx import Document
    from docx.shared import Inches
    doc=Document(); doc.add_heading("역학 분석 보고서",0)
    doc.add_paragraph(f"{fc.lab(outcome)} · {dataset} · 설계가중 로지스틱")
    doc.add_paragraph(f"[LOGISTIC] outcome = {outcome}")
    doc.add_heading("Figure 1. 연구대상자 선정 흐름도",1)
    doc.add_paragraph(f"전체 {n_total:,}명 → 연령 {amin}세 이상 {n_age:,}명 → 설계·결과 완전자료 {n_ana:,}명(분석대상).")
    def add_tab(title, df, note=""):
        doc.add_heading(title,1)
        if note: doc.add_paragraph(note)
        cols=list(df.columns); tb=doc.add_table(rows=1,cols=len(cols)); tb.style="Light Grid Accent 1"
        for j,c in enumerate(cols): tb.rows[0].cells[j].text=str(c)
        for _,row in df.iterrows():
            cc=tb.add_row().cells
            for j,c in enumerate(cols): cc[j].text=str(row[c])
    add_tab("Table 1. Baseline characteristics (weighted, with SMD)", t1)
    cr=crude.copy(); cr["OR (95% CI)"]=cr.apply(lambda x:f"{x.OR:.2f} ({x.lo:.2f}-{x.hi:.2f})" if pd.notna(x.OR) else "-",axis=1)
    cr["변수"]=cr.term.map(labels); cr["p"]=cr.p.map(lambda p:"<0.001" if pd.notna(p) and p<0.001 else (f"{p:.3f}" if pd.notna(p) else ""))
    add_tab("Table 2. Crude OR (univariable)", cr[["변수","OR (95% CI)","p"]])
    aj=adj.copy(); aj["OR (95% CI)"]=aj.apply(lambda x:f"{x.OR:.2f} ({x.lo:.2f}-{x.hi:.2f})" if pd.notna(x.OR) else "-",axis=1)
    aj["변수"]=aj.term.map(labels); aj["p"]=aj.p.map(lambda p:"<0.001" if pd.notna(p) and p<0.001 else (f"{p:.3f}" if pd.notna(p) else ""))
    add_tab("Table 3. Adjusted OR (multivariable)", aj[["변수","OR (95% CI)","p"]])
    add_tab("Table 4. 전체 위험 (weighted)", pd.DataFrame({"N":[n_ana],"Weighted risk (%)":[f"{risk.weighted_risk_pct[0]:.1f}"]}))
    SUBLAB={"men":"Sex","age50":"Age group"}
    def lvlab(s,lv):
        if s=="men": return "Men" if lv==1 else "Women"
        if s=="age50": return "≥50" if lv==1 else "<50"
        return str(lv)
    sg=sub.copy(); sg["OR (95% CI)"]=sg.apply(lambda x:f"{x.OR:.2f} ({x.lo:.2f}-{x.hi:.2f})" if pd.notna(x.OR) else "-",axis=1)
    sg["하위군"]=sg.subgroup.map(lambda s:SUBLAB.get(s, fc.lab(s) if s in fc.VARS else s))
    sg["수준"]=sg.apply(lambda r:lvlab(r.subgroup,r.level),axis=1)
    sg["P-interaction"]=sg.p_int.map(lambda p:"" if pd.isna(p) else ("<0.001" if p<0.001 else f"{p:.3f}"))
    add_tab(f"Table 5. 하위군 분석 — {fc.lab(main_exposure)} 효과 + P for interaction",
            sg[["하위군","수준","OR (95% CI)","P-interaction"]])
    doc.add_heading("Adjusted OR forest",1)
    bio=io.BytesIO(); fig.savefig(bio,format="png",dpi=110,bbox_inches="tight"); bio.seek(0)
    import matplotlib.pyplot as plt; plt.close(fig)
    doc.add_picture(bio,width=Inches(5.5))
    doc.add_heading("E-value",1)
    doc.add_paragraph(f"주요 노출 {fc.lab(main_exposure)} adjusted OR {me.OR:.2f}에 대한 E-value = {evp} (신뢰구간 한계 {evc}).")
    doc.add_heading("해석",1); doc.add_paragraph(interp)
    # ── Table 6: 노출효과 종합 (Crude·Min-adj·Full-adj·IPTW·G-comp·AIPW) ──
    doc.add_heading(f"Table 6. Main exposure 효과 종합 — {xbin_label}",1)
    t6=tab6.copy(); t6["Estimate (95% CI)"]=t6.apply(
        lambda x:f"{x.estimate:.2f} ({x.lo:.2f}-{x.hi:.2f})" if pd.notna(x.lo) else f"{x.estimate:.3f}",axis=1)
    add_tab("", t6[["method","scale","Estimate (95% CI)"]].rename(columns={"method":"방법","scale":"척도"}))
    # ── Love plot (IPTW 전후 SMD) ──
    doc.add_heading("Love plot — 공변량 균형 (IPTW 전후 |SMD|)",1)
    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
    lv=love.dropna(); figl,axl=plt.subplots(figsize=(6,0.4*len(lv)+1)); yy=np.arange(len(lv))
    axl.scatter(lv.before,yy,label="before",color="#A32D2D"); axl.scatter(lv.after,yy,label="after (IPTW)",color="#0F6E56")
    axl.axvline(0.1,color="#999",ls="--"); axl.set_yticks(yy); axl.set_yticklabels([labels.get(c,c) for c in lv.covariate],fontsize=8)
    axl.set_xlabel("|SMD|"); axl.legend(fontsize=8); axl.invert_yaxis(); figl.tight_layout()
    bl=io.BytesIO(); figl.savefig(bl,format="png",dpi=110,bbox_inches="tight"); bl.seek(0); plt.close(figl)
    doc.add_picture(bl,width=Inches(5.5))
    # ── VIF ──
    doc.add_heading("VIF (다중공선성)",1)
    vt=vif.copy(); vt["변수"]=vt.term.map(lambda t:labels.get(t,t)); vt["VIF"]=vt.VIF
    add_tab("", vt[["변수","VIF"]])
    # ── RCS 비선형성 ──
    pnl=rcs.value.iloc[0]
    doc.add_heading("제한 3차 스플라인 비선형성 검정",1)
    doc.add_paragraph("연속 노출에 미적용(이분 노출)." if pd.isna(pnl) else
        f"RCS(3 knots) 비선형성 P = {'<0.001' if pnl<0.001 else f'{pnl:.3f}'} "
        f"({'비선형성 유의' if pnl<0.05 else '선형성 기각 못함'}).")
    buf=io.BytesIO(); doc.save(buf); return buf.getvalue()
