# -*- coding: utf-8 -*-
"""발생/유병 트렌드 Word 보고서 (연도별 유병률·age-sex 표준화·Projection·층화·Joinpoint APC·음이항 예측)"""
import factory_core as fc, pandas as pd, numpy as np, subprocess, io, os, json

YMAP={"08":2008,"09":2009,"10":2010,"11":2011,"j":2018}
def _wprev(g,outcome,w):
    s=g[[outcome,w]].dropna();
    return 100*np.average(s[outcome].astype(float),weights=s[w]) if len(s) else np.nan

def build_trend_report(dataset, data_dir, years, outcome, defs, model, url, use_llm, workdir="."):
    amin=defs["pop"]["age_min"]; frames={}
    for y in years:
        DF=fc.load_raw(dataset,data_dir,(y,)); d=fc.apply_definitions(DF,defs)
        w="wt_tot" if "wt_tot" in d.columns else "wt_pool"
        dd=d[(d.age>=amin)&d[w].notna()&(d[w]>0)&d[outcome].notna()].copy()
        dd["_w"]=dd[w]; dd["ageg"]=pd.cut(dd.age,[amin,40,50,60,70,200],right=False,
                                          labels=["<40","40-49","50-59","60-69","70+"])
        dd["stratum"]=dd.ageg.astype(str)+"_"+dd.men.astype(str); frames[y]=dd
    alld=pd.concat(frames.values()); refw=alld.groupby("stratum")["_w"].sum(); refw/=refw.sum()
    rows=[]
    for y in years:
        dd=frames[y]; yr=YMAP.get(y,int("20"+y))
        sp=dd.groupby("stratum").apply(lambda g:_wprev(g,outcome,"_w"))
        std=np.nansum([refw.get(s,0)*sp.get(s,np.nan) for s in refw.index])
        rows.append({"year":yr,"N":len(dd),"count":int(round(dd[outcome].sum())),
                     "crude":_wprev(dd,outcome,"_w"),"standardized":std,
                     "men":_wprev(dd[dd.men==1],outcome,"_w"),"women":_wprev(dd[dd.men==0],outcome,"_w")})
    T=pd.DataFrame(rows).sort_values("year").reset_index(drop=True)
    # 선형 Projection (표준화율 기준, +3년, 회귀 예측구간)
    x=T.year.values; yme=T.standardized.values; n=len(x)
    b1,b0=np.polyfit(x,yme,1); yhat=b0+b1*x; sse=np.sum((yme-yhat)**2); se=np.sqrt(sse/max(n-2,1))
    fut=np.arange(x.max()+1,x.max()+4)
    xm=x.mean(); sxx=np.sum((x-xm)**2)
    proj=[]
    for xf in fut:
        yf=b0+b1*xf; pm=se*np.sqrt(1+1/n+(xf-xm)**2/sxx)
        proj.append({"year":int(xf),"forecast":round(yf,2),"lo95":round(yf-1.96*pm,2),"hi95":round(yf+1.96*pm,2)})
    proj=pd.DataFrame(proj)
    # trend.R (NB 예측 + APC 부트스트랩)
    T[["year","count","N","standardized"]].rename(columns={"standardized":"rate"}).to_csv(os.path.join(workdir,"trend_series.csv"),index=False)
    json.dump({"future":fut.tolist()},open(os.path.join(workdir,"trend_config.json"),"w"))
    r=subprocess.run([fc.rscript_cmd(),"trend.R"],cwd=workdir,capture_output=True,text=True,env=fc.r_env(workdir))
    if r.returncode!=0: raise RuntimeError(r.stderr)
    nb=pd.read_csv(os.path.join(workdir,"trend_nb.csv")); apc=pd.read_csv(os.path.join(workdir,"trend_apc.csv"))
    # 그림
    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
    fig,ax=plt.subplots(figsize=(7,4))
    ax.plot(T.year,T.crude,"o-",label="Crude",color="#245")
    ax.plot(T.year,T.standardized,"s-",label="Age-sex standardized",color="#0F6E56")
    ax.plot(proj.year,proj.forecast,"^--",label="Projection",color="#A32D2D")
    ax.fill_between(proj.year,proj.lo95,proj.hi95,color="#A32D2D",alpha=0.15)
    ax.set_xlabel("Year"); ax.set_ylabel(f"{fc.lab(outcome)} prevalence, %"); ax.legend(fontsize=8); ax.grid(alpha=.3)
    fig.tight_layout(); bio=io.BytesIO(); fig.savefig(bio,format="png",dpi=110,bbox_inches="tight"); bio.seek(0); plt.close(fig)
    # 해석
    a=apc.iloc[0]; facts=f"{fc.lab(outcome)} annual percent change {a.APC_pct}% (95% CI {a.lo95} to {a.hi95})."
    interp=""
    if use_llm:
        try: interp=fc.ollama_chat(f"{fc.STYLE}\n추세분석 결과를 논문 문장으로. 수치변경 금지.\n{facts}",model,url).strip()
        except Exception: interp=""
    if not interp:
        d_="증가" if a.APC_pct>0 else "감소"
        interp=(f"The age-sex standardized prevalence of {fc.lab(outcome)} showed an annual percent change of "
                f"{a.APC_pct}% (95% CI {a.lo95} to {a.hi95}) over the study period.")
    # ── docx ──
    from docx import Document; from docx.shared import Inches
    doc=Document(); doc.add_heading("발생/유병 트렌드 보고서",0)
    doc.add_paragraph(f"{fc.lab(outcome)} · {dataset} · {int(T.year.min())}~{int(T.year.max())} · 설계가중")
    doc.add_heading("Figure. 발생/유병 추세",1)
    doc.add_picture(bio,width=Inches(6.0))
    def add_tab(title,df):
        doc.add_heading(title,1); cols=list(df.columns); tb=doc.add_table(rows=1,cols=len(cols)); tb.style="Light Grid Accent 1"
        for j,c in enumerate(cols): tb.rows[0].cells[j].text=str(c)
        for _,row in df.iterrows():
            cc=tb.add_row().cells
            for j,c in enumerate(cols): cc[j].text=(f"{row[c]:.2f}" if isinstance(row[c],float) else str(row[c]))
    add_tab("시간단위 건수 (period, count)", T[["year","count"]].rename(columns={"year":"period"}))
    add_tab("Crude · age-sex 표준화율 (%)", T[["year","crude","standardized"]])
    add_tab("Projection (선형 예측, 표준화율)", proj)
    add_tab("층화 추세 — 성별 (%)", T[["year","men","women"]])
    add_tab("Joinpoint 회귀 — APC (%) + 부트스트랩 95% CI", apc.rename(columns={"APC_pct":"APC(%)"}))
    add_tab("음이항(NB) 예측 (%)", nb)
    doc.add_heading("해석",1); doc.add_paragraph(interp)
    if len(years)<5: doc.add_paragraph(f"주: 현재 {len(years)}개 시점만 사용됨. 사이클을 더 추가하면 추세·Joinpoint 추정이 안정화됩니다.")
    buf=io.BytesIO(); doc.save(buf); return buf.getvalue()
