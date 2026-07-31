# -*- coding: utf-8 -*-
"""ML prediction Word report -- tuning, search space, full metrics, ROC/PR, threshold, calibration, SHAP, PDP"""
import factory_core as fc, pandas as pd, numpy as np, io, warnings
warnings.filterwarnings("ignore")
from sklearn.model_selection import RandomizedSearchCV, train_test_split, StratifiedKFold
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier, ExtraTreesClassifier
from sklearn.neighbors import KNeighborsClassifier
from sklearn.impute import SimpleImputer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import (roc_auc_score, average_precision_score, confusion_matrix, f1_score,
    matthews_corrcoef, cohen_kappa_score, log_loss, brier_score_loss, roc_curve, precision_recall_curve)
from sklearn.calibration import calibration_curve
from xgboost import XGBClassifier

SPACES={
 "LogisticRegression":(LogisticRegression(max_iter=2000,solver="liblinear"),
    {"clf__C":[0.01,0.1,1,10],"clf__penalty":["l1","l2"]}),
 "RandomForest":(RandomForestClassifier(random_state=0),
    {"clf__n_estimators":[200,400],"clf__max_depth":[4,6,8,None],"clf__min_samples_leaf":[1,5,10]}),
 "ExtraTrees":(ExtraTreesClassifier(random_state=0),
    {"clf__n_estimators":[200,400],"clf__max_depth":[6,10,None],"clf__min_samples_leaf":[1,5]}),
 "GradientBoosting":(GradientBoostingClassifier(random_state=0),
    {"clf__n_estimators":[150,300],"clf__learning_rate":[0.03,0.1],"clf__max_depth":[2,3]}),
 "XGBoost":(XGBClassifier(random_state=0,eval_metric="logloss",verbosity=0),
    {"clf__n_estimators":[200,400],"clf__max_depth":[3,5],"clf__learning_rate":[0.03,0.1],"clf__subsample":[0.8,1.0]}),
 "KNN":(KNeighborsClassifier(),{"clf__n_neighbors":[5,15,31],"clf__weights":["uniform","distance"]}),
}

def _metrics(y,p,thr=0.5):
    pred=(p>=thr).astype(int); tn,fp,fn,tp=confusion_matrix(y,pred,labels=[0,1]).ravel()
    sens=tp/(tp+fn) if tp+fn else np.nan; spec=tn/(tn+fp) if tn+fp else np.nan
    ppv=tp/(tp+fp) if tp+fp else np.nan; npv=tn/(tn+fn) if tn+fn else np.nan
    return dict(TP=tp,TN=tn,FP=fp,FN=fn,sensitivity=sens,specificity=spec,PPV=ppv,NPV=npv,
        accuracy=(tp+tn)/len(y),balanced_accuracy=np.nanmean([sens,spec]),
        f1=f1_score(y,pred,zero_division=0),MCC=matthews_corrcoef(y,pred),kappa=cohen_kappa_score(y,pred),
        auROC=roc_auc_score(y,p),auPRC=average_precision_score(y,p),
        logloss=log_loss(y,np.clip(p,1e-6,1-1e-6)),brier=brier_score_loss(y,p))

def build_ml_report(dataset, DF, outcome, predictors, defs, model, url, use_llm, workdir="."):
    d=fc.apply_definitions(DF,defs); amin=defs["pop"]["age_min"]
    feats=[f for f in predictors if f in d.columns and f!=outcome]
    ana=d[(d.age>=amin)&d[outcome].notna()].copy()
    X=ana[feats].apply(pd.to_numeric,errors="coerce"); y=ana[outcome].astype(int).values
    keep=X.notna().mean(axis=0)>0.3; X=X.loc[:,keep]; feats=list(X.columns)
    Xtr,Xte,ytr,yte=train_test_split(X.values,y,test_size=0.3,random_state=0,stratify=y)
    cv=StratifiedKFold(3,shuffle=True,random_state=0)
    tuned={}; tune_rows=[]
    for name,(est,grid) in SPACES.items():
        pipe=Pipeline([("imp",SimpleImputer(strategy="median")),("sc",StandardScaler()),("clf",est)])
        rs=RandomizedSearchCV(pipe,grid,n_iter=min(8,int(np.prod([len(v) for v in grid.values()]))),
            scoring="roc_auc",cv=cv,random_state=0,n_jobs=-1)
        rs.fit(Xtr,ytr); tuned[name]=rs.best_estimator_
        bp={k.replace("clf__",""):v for k,v in rs.best_params_.items()}
        tune_rows.append({"model":name,"n_iter":rs.n_iter_ if hasattr(rs,"n_iter_") else len(rs.cv_results_["params"]),
                          "CV_auROC":round(rs.best_score_,4),"best_params":str(bp)})
    # Full performance metrics (test)
    rows=[]; probs={}
    for name,estm in tuned.items():
        p=estm.predict_proba(Xte)[:,1]; probs[name]=p; m=_metrics(yte,p)
        m["combined"]=round((m["auROC"]+m["auPRC"])/2,4); m={"model":name,**{k:(round(v,4) if isinstance(v,float) else v) for k,v in m.items()}}
        rows.append(m)
    metrics=pd.DataFrame(rows).sort_values("combined",ascending=False).reset_index(drop=True)
    best=metrics.model.iloc[0]; best_est=tuned[best]; pbest=probs[best]
    # Curve figures
    import matplotlib; matplotlib.use("Agg"); import matplotlib.pyplot as plt
    def fig_roc_pr():
        fig,ax=plt.subplots(1,2,figsize=(9,4))
        for name,p in probs.items():
            fpr,tpr,_=roc_curve(yte,p); ax[0].plot(fpr,tpr,lw=1,label=f"{name} ({roc_auc_score(yte,p):.3f})")
            pr,rc,_=precision_recall_curve(yte,p); ax[1].plot(rc,pr,lw=1,label=name)
        ax[0].plot([0,1],[0,1],"--",color="#999"); ax[0].set_title("ROC"); ax[0].set_xlabel("FPR"); ax[0].set_ylabel("TPR"); ax[0].legend(fontsize=6)
        ax[1].set_title("Precision-Recall"); ax[1].set_xlabel("Recall"); ax[1].set_ylabel("Precision")
        fig.tight_layout(); return fig
    def fig_cal():
        fig,ax=plt.subplots(figsize=(5,4)); fr,mp=calibration_curve(yte,pbest,n_bins=10,strategy="quantile")
        ax.plot(mp,fr,"o-",color="#0F6E56"); ax.plot([0,1],[0,1],"--",color="#999")
        ax.set_title(f"Calibration -- {best}"); ax.set_xlabel("Predicted"); ax.set_ylabel("Observed"); fig.tight_layout(); return fig
    # Threshold analysis
    thr_rows=[]
    for t in np.arange(0.1,0.95,0.1):
        m=_metrics(yte,pbest,t); thr_rows.append({"threshold":round(t,1),**{k:m[k] for k in ["TP","TN","FP","FN","sensitivity","specificity","PPV","NPV","f1","accuracy"]}})
    thr=pd.DataFrame(thr_rows)
    for c in ["sensitivity","specificity","PPV","NPV","f1","accuracy"]: thr[c]=thr[c].round(3)
    # SHAP (TreeExplainer for tree models, otherwise importance fallback)
    import shap
    Xte_df=pd.DataFrame(Xte,columns=feats); samp=Xte_df.sample(min(150,len(Xte_df)),random_state=0)
    shap_fig=None
    try:
        clf=best_est.named_steps["clf"]; pre=Pipeline(best_est.steps[:-1])
        Xs=pre.transform(samp.values)
        if best in ("RandomForest","ExtraTrees","GradientBoosting","XGBoost"):
            ex=shap.TreeExplainer(clf); sv=ex.shap_values(Xs)
            sv=sv[1] if isinstance(sv,list) else sv
        else:
            ex=shap.LinearExplainer(clf,Xs); sv=ex.shap_values(Xs)
        plt.figure(); shap.summary_plot(sv,pd.DataFrame(Xs,columns=feats),show=False,plot_size=(6,4))
        shap_fig=plt.gcf()
    except Exception as e:
        imp=pd.Series(getattr(best_est.named_steps["clf"],"feature_importances_",np.zeros(len(feats))),index=feats).sort_values()
        fig,ax=plt.subplots(figsize=(6,4)); ax.barh(imp.index,imp.values,color="#245"); ax.set_title("Feature importance"); fig.tight_layout(); shap_fig=fig
    # PDP (top 2 features)
    from sklearn.inspection import PartialDependenceDisplay
    try:
        imp=getattr(best_est.named_steps["clf"],"feature_importances_",None)
        top=list(np.argsort(imp)[::-1][:2]) if imp is not None else [0,1]
        figp,axp=plt.subplots(1,len(top),figsize=(4*len(top),3.5))
        PartialDependenceDisplay.from_estimator(best_est,Xte,top,feature_names=feats,ax=axp); figp.tight_layout()
    except Exception:
        figp,axp=plt.subplots(figsize=(4,3)); axp.text(0.5,0.5,"PDP unavailable",ha="center"); figp.tight_layout()
    # Interpretation
    bm=metrics.iloc[0]; facts=f"Best model {best}: auROC {bm.auROC}, auPRC {bm.auPRC}, combined {bm.combined}."
    interp=""
    if use_llm:
        try: interp=fc.ollama_chat(f"{fc.STYLE}\nWrite the ML prediction-model results below as journal prose. Do not change any numbers.\n{facts}",model,url).strip()
        except Exception: interp=""
    if not interp:
        interp=(f"Among the evaluated models, {best} achieved the highest combined score, with an area under the "
                f"receiver operating characteristic curve of {bm.auROC} and an area under the precision-recall curve of {bm.auPRC}.")
    # -- docx --
    from docx import Document; from docx.shared import Inches
    def emb(fig,w=6.0):
        b=io.BytesIO(); fig.savefig(b,format="png",dpi=110,bbox_inches="tight"); b.seek(0); plt.close(fig); doc.add_picture(b,width=Inches(w))
    def tab(title,df,note=""):
        doc.add_heading(title,1)
        if note: doc.add_paragraph(note)
        cols=list(df.columns); tb=doc.add_table(rows=1,cols=len(cols)); tb.style="Light Grid Accent 1"
        for j,c in enumerate(cols): tb.rows[0].cells[j].text=str(c)
        for _,r in df.iterrows():
            cc=tb.add_row().cells
            for j,c in enumerate(cols): cc[j].text=str(r[c])
    doc=Document(); doc.add_heading("ML Prediction Model Report",0)
    doc.add_paragraph(f"{fc.lab(outcome)} - {dataset} - search space, tuning, full metrics, threshold, SHAP, PDP")
    doc.add_paragraph(f"[ML] outcome = {outcome}")
    tab("Data summary", pd.DataFrame({"Item":["N","Outcome prevalence (%)","Number of features","Train/test"],
        "Value":[len(ana),f"{100*y.mean():.1f}",len(feats),f"{len(ytr)}/{len(yte)} (70/30 stratified)"]}))
    ss=pd.DataFrame([{"Model":k,"Search space":str({kk.replace('clf__',''):vv for kk,vv in g.items()})} for k,(e,g) in SPACES.items()])
    tab("Hyperparameter search space", ss)
    tab("Tuning results -- best params, CV auROC", pd.DataFrame(tune_rows))
    tab("Full performance metrics (test set)", metrics)
    doc.add_heading("ROC / Precision-Recall",1); emb(fig_roc_pr(),6.5)
    doc.add_heading(f"Final selected model -- {best} (combined={bm.combined})",1); doc.add_paragraph(interp)
    tab("Threshold analysis (0.1-0.9)", thr)
    doc.add_heading("Calibration",1); emb(fig_cal(),4.8)
    doc.add_heading("SHAP feature contributions",1); emb(shap_fig,6.0)
    doc.add_heading("Partial dependence (PDP)",1); emb(figp,4*len(feats[:2]) if feats else 4)
    buf=io.BytesIO(); doc.save(buf); return buf.getvalue()
